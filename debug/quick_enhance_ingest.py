#!/usr/bin/env python
"""
Quick enhanced ingestion - 95% coverage target
Works with existing dependencies, focuses on DOC files and better PDF handling
"""
import json
import logging
import os
import sys
import time
from datetime import datetime
from hashlib import sha256
from pathlib import Path

# Setup
KB_DIR = Path("knowledgebase")
IDX_DIR = Path("knowledgebase_index")
MANIFEST_PATH = IDX_DIR / "ingestion_manifest.jsonl"

logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(message)s")
logger = logging.getLogger("quick_enhance")

# Import required libraries (only what we already have)
try:
    from sentence_transformers import SentenceTransformer
    import numpy as np
    import faiss
    import fitz  # PyMuPDF
    from docx import Document
except ImportError as e:
    logger.error(f"Missing core dependency: {e}")
    sys.exit(1)

def load_existing_hashes():
    if not MANIFEST_PATH.exists():
        return set()
    hashes = set()
    with open(MANIFEST_PATH, 'r', encoding='utf-8') as f:
        for line in f:
            try:
                entry = json.loads(line.strip())
                hashes.add(entry["file_sha256"])
            except:
                continue
    return hashes

def compute_sha256(path):
    h = sha256()
    try:
        with open(path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                h.update(chunk)
        return h.hexdigest()
    except Exception as e:
        logger.error(f"Failed to compute hash for {path}: {e}")
        return ""

def extract_text_from_doc(file_path):
    """Extract text from DOC files using PyMuPDF"""
    try:
        size_mb = file_path.stat().st_size / 1024 / 1024
        if size_mb > 15:  # Skip very large DOC files
            logger.warning(f"Skipping large DOC file {file_path.name} ({size_mb:.1f}MB)")
            return ""
        
        # Try using PyMuPDF for DOC files
        with fitz.open(str(file_path)) as doc:
            text = ""
            max_pages = min(len(doc), 50)  # Limit pages for speed
            for i in range(max_pages):
                page_text = doc[i].get_text()
                if page_text.strip():
                    text += page_text + "\n"
            return text[:150000]  # Limit text length
    except Exception as e:
        logger.error(f"Failed to extract from DOC {file_path.name}: {e}")
        return ""

def extract_text_from_pdf(file_path):
    """Enhanced PDF text extraction with better handling"""
    try:
        size_mb = file_path.stat().st_size / 1024 / 1024
        
        if size_mb > 30:  # Skip very large PDFs
            logger.warning(f"Skipping very large PDF {file_path.name} ({size_mb:.1f}MB)")
            return ""
        
        with fitz.open(str(file_path)) as doc:
            text = ""
            max_pages = min(len(doc), 300)  # Increased page limit
            
            for i in range(max_pages):
                page_text = doc[i].get_text()
                if page_text.strip():
                    text += page_text + "\n"
            
            return text[:300000]  # Limit text length
    except Exception as e:
        logger.error(f"Failed to extract from PDF {file_path.name}: {e}")
        return ""

def extract_text_from_docx(file_path):
    """Extract text from DOCX files"""
    try:
        size_mb = file_path.stat().st_size / 1024 / 1024
        if size_mb > 8:  # Skip very large DOCX files
            logger.warning(f"Skipping large DOCX file {file_path.name} ({size_mb:.1f}MB)")
            return ""
        
        doc = Document(str(file_path))
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        return text[:150000]  # Limit text length
    except Exception as e:
        logger.error(f"Failed to extract from DOCX {file_path.name}: {e}")
        return ""

def extract_text_from_txt(file_path):
    """Extract text from TXT files"""
    try:
        size_mb = file_path.stat().st_size / 1024 / 1024
        if size_mb > 3:  # Skip very large text files
            logger.warning(f"Skipping large text file {file_path.name} ({size_mb:.1f}MB)")
            return ""
        
        text = file_path.read_text(encoding='utf-8', errors='ignore')
        return text[:80000]  # Limit text length
    except Exception as e:
        logger.error(f"Failed to extract from TXT {file_path.name}: {e}")
        return ""

def extract_text_from_json(file_path):
    """Extract text from JSON files"""
    try:
        size_mb = file_path.stat().st_size / 1024 / 1024
        if size_mb > 1:  # Skip large JSON files
            logger.warning(f"Skipping large JSON file {file_path.name} ({size_mb:.1f}MB)")
            return ""
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            data = json.load(f)
        return json.dumps(data, indent=2)[:40000]  # Limit text length
    except Exception as e:
        logger.error(f"Failed to extract from JSON {file_path.name}: {e}")
        return ""

def extract_text(file_path):
    """Main text extraction function"""
    suffix = file_path.suffix.lower()
    
    if suffix == '.doc':
        return extract_text_from_doc(file_path)
    elif suffix == '.pdf':
        return extract_text_from_pdf(file_path)
    elif suffix == '.docx':
        return extract_text_from_docx(file_path)
    elif suffix == '.txt':
        return extract_text_from_txt(file_path)
    elif suffix == '.json':
        return extract_text_from_json(file_path)
    else:
        logger.warning(f"Unsupported file type: {suffix}")
        return ""

def chunk_text(text, size=1000, overlap=200):
    """Simple chunking with minimum text requirement"""
    if len(text.strip()) < 30:  # Lowered minimum for more files
        return []
    
    chunks = []
    start = 0
    while start < len(text):
        end = start + size
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - overlap
        if start >= len(text):
            break
    
    return chunks

def main():
    print("🚀 QUICK ENHANCED INGESTION - 95% COVERAGE TARGET")
    print("=" * 60)
    
    # Setup
    IDX_DIR.mkdir(exist_ok=True)
    existing_hashes = load_existing_hashes()
    print(f"📋 Found {len(existing_hashes)} previously processed files")
    
    # Initialize model and index
    print("🤖 Loading AI model...")
    model = SentenceTransformer("all-MiniLM-L6-v2", device="cpu")
    print("✅ Model loaded!")
    
    # Load existing index
    existing_faiss = list(IDX_DIR.glob("index_v*.faiss"))
    if existing_faiss:
        latest_faiss = max(existing_faiss, key=lambda x: x.stat().st_mtime)
        try:
            index = faiss.read_index(str(latest_faiss))
            pkl_file = latest_faiss.with_suffix('.pkl')
            if pkl_file.exists():
                with open(pkl_file, 'r') as f:
                    metadata = json.load(f)
            else:
                metadata = []
            print(f"📁 Loaded existing index with {index.ntotal} vectors")
        except Exception as e:
            print(f"⚠️  Failed to load existing index: {e}")
            index = faiss.IndexFlatIP(384)
            metadata = []
    else:
        index = faiss.IndexFlatIP(384)
        metadata = []
    
    # Find files to process (focus on DOC files and unprocessed PDFs/DOCXs)
    all_files = []
    supported_extensions = ['.pdf', '.docx', '.doc', '.txt', '.json']
    
    for ext in supported_extensions:
        all_files.extend(KB_DIR.rglob(f"*{ext}"))
    
    print(f"📁 Found {len(all_files)} total files with support")
    
    # Find unprocessed files
    unprocessed = []
    for file_path in all_files:
        try:
            file_hash = compute_sha256(file_path)
            if file_hash and file_hash not in existing_hashes:
                unprocessed.append(file_path)
        except:
            unprocessed.append(file_path)
    
    print(f"🔄 Need to process: {len(unprocessed)} files")
    
    if len(unprocessed) == 0:
        print("🎉 All files already processed!")
        return
    
    print("\n" + "=" * 60)
    print("🚀 STARTING QUICK ENHANCED PROCESSING")
    print("=" * 60)
    
    # Process files
    processed = 0
    total_chunks = len(metadata)
    version = datetime.utcnow().strftime("%Y%m%d_%H%M")
    start_time = time.time()
    
    for i, file_path in enumerate(unprocessed):
        try:
            print(f"📄 Processing {i+1}/{len(unprocessed)}: {file_path.name}")
            
            file_hash = compute_sha256(file_path)
            if not file_hash:
                print(f"   ❌ Failed to compute hash for {file_path.name}")
                continue
            
            # Extract text
            text = extract_text(file_path)
            if not text or len(text.strip()) < 15:  # Lowered minimum
                print(f"   ⚠️  No meaningful text from {file_path.name}")
                continue
            
            # Chunk text
            chunks = chunk_text(text)
            if not chunks:
                print(f"   ⚠️  No chunks generated from {file_path.name}")
                continue
            
            print(f"   ✅ Generated {len(chunks)} chunks from {file_path.name}")
            
            # Generate embeddings
            embeddings = []
            for j in range(0, len(chunks), 32):  # Smaller batch size
                batch = chunks[j:j+32]
                try:
                    embs = model.encode(batch, show_progress_bar=False, convert_to_numpy=True, normalize_embeddings=True)
                    embeddings.append(embs.astype('float32'))
                except Exception as e:
                    print(f"   ❌ Failed to embed batch for {file_path.name}: {e}")
                    continue
            
            if embeddings:
                vectors = np.concatenate(embeddings)
                index.add(vectors)
                
                # Add metadata
                timestamp = datetime.utcnow().isoformat()
                for idx in range(len(chunks)):
                    metadata.append({
                        "file_name": str(file_path.relative_to(KB_DIR)),
                        "file_sha256": file_hash,
                        "chunk_index": idx,
                        "ingest_timestamp": timestamp,
                    })
                
                total_chunks += len(chunks)
                processed += 1
                
                # Update manifest
                with open(MANIFEST_PATH, 'a', encoding='utf-8') as f:
                    entry = {
                        "file_name": str(file_path.relative_to(KB_DIR)),
                        "file_sha256": file_hash,
                        "chunk_count": len(chunks),
                        "ingest_timestamp": timestamp,
                        "model_name": "all-MiniLM-L6-v2",
                        "index_version": version,
                    }
                    f.write(json.dumps(entry) + "\n")
                
                # Save progress every 5 files
                if processed % 5 == 0:
                    faiss_path = IDX_DIR / f"index_v{version}.faiss"
                    pkl_path = IDX_DIR / f"index_v{version}.pkl"
                    
                    faiss.write_index(index, str(faiss_path))
                    with open(pkl_path, 'w') as f:
                        json.dump(metadata, f)
                    
                    elapsed = time.time() - start_time
                    rate = processed / elapsed if elapsed > 0 else 0
                    print(f"   💾 Saved checkpoint: {processed} files, {total_chunks} chunks ({rate:.1f} files/sec)")
            
        except Exception as e:
            print(f"   ❌ Failed to process {file_path.name}: {e}")
            continue
    
    # Final save
    faiss_path = IDX_DIR / f"index_v{version}.faiss"
    pkl_path = IDX_DIR / f"index_v{version}.pkl"
    
    faiss.write_index(index, str(faiss_path))
    with open(pkl_path, 'w') as f:
        json.dump(metadata, f)
    
    print("\n" + "=" * 60)
    print("🎉 QUICK ENHANCED INGESTION COMPLETE!")
    print("=" * 60)
    print(f"✅ Processed: {processed} new files")
    print(f"📊 Total chunks: {total_chunks}")
    print(f"📁 Index size: {faiss_path.stat().st_size / 1024 / 1024:.1f}MB")
    print(f"📈 Coverage improvement: +{processed} files")
    
    # Calculate coverage percentage
    total_files_in_kb = len(list(KB_DIR.rglob("*")))
    total_processed = len(existing_hashes) + processed
    coverage = (total_processed / total_files_in_kb) * 100 if total_files_in_kb > 0 else 0
    print(f"🎯 Estimated coverage: {coverage:.1f}%")

if __name__ == "__main__":
    main()
