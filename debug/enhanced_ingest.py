#!/usr/bin/env python
"""
Enhanced knowledgebase ingestion - 95% coverage target
Processes images, DOC files, and handles large files better
"""
import json
import logging
import os
import sys
import time
from datetime import datetime
from hashlib import sha256
from pathlib import Path
from typing import List, Optional

# Setup
KB_DIR = Path("knowledgebase")
CONVERTED_DIR = KB_DIR / "converted_docs"
IDX_DIR = Path("knowledgebase_index")
MANIFEST_PATH = IDX_DIR / "ingestion_manifest.jsonl"

logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(message)s")
logger = logging.getLogger("enhanced_ingest")

# Import required libraries
try:
    from sentence_transformers import SentenceTransformer
    import numpy as np
    import faiss
    import fitz  # PyMuPDF
    from docx import Document
    from PIL import Image, ImageFile
    import pytesseract
    import pandas as pd
except ImportError as e:
    logger.error(f"Missing dependency: {e}")
    logger.info("Installing missing dependencies...")
    os.system("pip install sentence-transformers faiss-cpu PyMuPDF python-docx Pillow pytesseract pandas")
    sys.exit(1)

# Allow loading truncated JPEGs to avoid 'Unsupported image format/type'
try:
    ImageFile.LOAD_TRUNCATED_IMAGES = True
except Exception:
    pass

def load_existing_hashes():
    if not MANIFEST_PATH.exists():
        return set()
    hashes = set()
    with open(MANIFEST_PATH, 'r', encoding='utf-8') as f:
        for line in f:
            try:
                entry = json.loads(line.strip())
                hashes.add(entry["file_sha256"])
            except:
                continue
    return hashes

def compute_sha256(path):
    h = sha256()
    try:
        with open(path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                h.update(chunk)
        return h.hexdigest()
    except Exception as e:
        logger.error(f"Failed to compute hash for {path}: {e}")
        return ""

def extract_text_from_image(file_path):
    """Extract text from images using OCR"""
    try:
        size_mb = file_path.stat().st_size / 1024 / 1024
        if size_mb > 10:  # Skip very large images
            logger.warning(f"Skipping large image {file_path.name} ({size_mb:.1f}MB)")
            return ""
        
        image = Image.open(file_path)
        # Convert to RGB if needed
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Use OCR to extract text
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        logger.error(f"Failed to extract from image {file_path.name}: {e}")
        return ""

def extract_text_from_doc(file_path):
    """Legacy DOC handling: prefer converted PDF if available, otherwise skip."""
    try:
        rel = file_path.relative_to(KB_DIR)
        pdf_path = (CONVERTED_DIR / rel).with_suffix(".pdf")
        if pdf_path.exists():
            return extract_text_from_pdf(pdf_path)
        else:
            logger.warning(f"No converted PDF for DOC {file_path.name}; skipping")
            return ""
    except Exception as e:
        logger.error(f"DOC handling error for {file_path.name}: {e}")
        return ""

def extract_text_from_pdf(file_path):
    """Enhanced PDF text extraction"""
    try:
        size_mb = file_path.stat().st_size / 1024 / 1024
        
        if size_mb > 50:  # Skip extremely large PDFs
            logger.warning(f"Skipping very large PDF {file_path.name} ({size_mb:.1f}MB)")
            return ""
        
        with fitz.open(str(file_path)) as doc:
            text = ""
            max_pages = min(len(doc), 500)  # Increased page limit
            
            for i in range(max_pages):
                page_text = doc[i].get_text()
                if page_text.strip():
                    text += page_text + "\n"
            
            # If no text found, try OCR on first few pages
            if len(text.strip()) < 100 and size_mb < 10:
                logger.info(f"Attempting OCR on {file_path.name}")
                for i in range(min(5, len(doc))):
                    pix = doc[i].get_pixmap()
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ocr_text = pytesseract.image_to_string(img)
                    text += ocr_text + "\n"
            
            return text[:500000]  # Limit text length
    except Exception as e:
        logger.error(f"Failed to extract from PDF {file_path.name}: {e}")
        return ""

def extract_text_from_docx(file_path):
    """Extract text from DOCX files"""
    try:
        size_mb = file_path.stat().st_size / 1024 / 1024
        if size_mb > 10:  # Skip very large DOCX files
            logger.warning(f"Skipping large DOCX file {file_path.name} ({size_mb:.1f}MB)")
            return ""
        
        doc = Document(str(file_path))
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        return text[:200000]  # Limit text length
    except Exception as e:
        logger.error(f"Failed to extract from DOCX {file_path.name}: {e}")
        return ""

def extract_text_from_txt(file_path):
    """Extract text from TXT files"""
    try:
        size_mb = file_path.stat().st_size / 1024 / 1024
        if size_mb > 5:  # Skip very large text files
            logger.warning(f"Skipping large text file {file_path.name} ({size_mb:.1f}MB)")
            return ""
        
        text = file_path.read_text(encoding='utf-8', errors='ignore')
        return text[:100000]  # Limit text length
    except Exception as e:
        logger.error(f"Failed to extract from TXT {file_path.name}: {e}")
        return ""

def extract_text_from_json(file_path):
    """Extract text from JSON files"""
    try:
        size_mb = file_path.stat().st_size / 1024 / 1024
        if size_mb > 2:  # Skip large JSON files
            logger.warning(f"Skipping large JSON file {file_path.name} ({size_mb:.1f}MB)")
            return ""
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            data = json.load(f)
        return json.dumps(data, indent=2)[:50000]  # Limit text length
    except Exception as e:
        logger.error(f"Failed to extract from JSON {file_path.name}: {e}")
        return ""

def extract_text_from_csv(file_path):
    """Extract text from CSV files"""
    try:
        size_mb = file_path.stat().st_size / 1024 / 1024
        if size_mb > 5:  # Skip large CSV files
            logger.warning(f"Skipping large CSV file {file_path.name} ({size_mb:.1f}MB)")
            return ""
        
        df = pd.read_csv(str(file_path), dtype=str, low_memory=False)
        return df.to_string(index=False)[:50000]  # Limit text length
    except Exception as e:
        logger.error(f"Failed to extract from CSV {file_path.name}: {e}")
        return ""

def extract_text(file_path):
    """Main text extraction function with enhanced file type support"""
    suffix = file_path.suffix.lower()
    
    if suffix in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif']:
        return extract_text_from_image(file_path)
    elif suffix == '.doc':
        return extract_text_from_doc(file_path)
    elif suffix == '.pdf':
        return extract_text_from_pdf(file_path)
    elif suffix == '.docx':
        return extract_text_from_docx(file_path)
    elif suffix == '.txt':
        return extract_text_from_txt(file_path)
    elif suffix == '.json':
        return extract_text_from_json(file_path)
    elif suffix == '.csv':
        return extract_text_from_csv(file_path)
    else:
        logger.warning(f"Unsupported file type: {suffix}")
        return ""

def chunk_text(text, size=1000, overlap=200):
    """Simple chunking with minimum text requirement"""
    if len(text.strip()) < 50:
        return []
    
    chunks = []
    start = 0
    while start < len(text):
        end = start + size
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - overlap
        if start >= len(text):
            break
    
    return chunks

def main():
    print("🚀 ENHANCED KNOWLEDGEBASE INGESTION - 95% COVERAGE TARGET")
    print("=" * 70)
    
    # Setup
    IDX_DIR.mkdir(exist_ok=True)
    existing_hashes = load_existing_hashes()
    print(f"📋 Found {len(existing_hashes)} previously processed files")
    
    # Initialize model and index
    print("🤖 Loading AI model...")
    model = SentenceTransformer("all-MiniLM-L6-v2", device="cpu")
    print("✅ Model loaded!")
    
    # Load existing index
    existing_faiss = list(IDX_DIR.glob("index_v*.faiss"))
    if existing_faiss:
        latest_faiss = max(existing_faiss, key=lambda x: x.stat().st_mtime)
        try:
            index = faiss.read_index(str(latest_faiss))
            pkl_file = latest_faiss.with_suffix('.pkl')
            if pkl_file.exists():
                with open(pkl_file, 'r') as f:
                    metadata = json.load(f)
            else:
                metadata = []
            print(f"📁 Loaded existing index with {index.ntotal} vectors")
        except Exception as e:
            print(f"⚠️  Failed to load existing index: {e}")
            index = faiss.IndexFlatIP(384)
            metadata = []
    else:
        index = faiss.IndexFlatIP(384)
        metadata = []
    
    # Find all files with enhanced support
    all_files = []
    supported_extensions = [
        '.pdf', '.docx', '.txt', '.json', '.csv',
        '.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif'
    ]
    
    for ext in supported_extensions:
        all_files.extend(KB_DIR.rglob(f"*{ext}"))
    
    print(f"📁 Found {len(all_files)} total files with enhanced support")
    
    # Find unprocessed files
    unprocessed = []
    for file_path in all_files:
        try:
            file_hash = compute_sha256(file_path)
            if file_hash and file_hash not in existing_hashes:
                unprocessed.append(file_path)
        except:
            unprocessed.append(file_path)
    
    print(f"🔄 Need to process: {len(unprocessed)} files")
    
    if len(unprocessed) == 0:
        print("🎉 All files already processed!")
        return
    
    print("\n" + "=" * 70)
    print("🚀 STARTING ENHANCED PROCESSING")
    print("=" * 70)
    
    # Process files
    processed = 0
    total_chunks = len(metadata)
    version = datetime.utcnow().strftime("%Y%m%d_%H%M")
    start_time = time.time()
    
    for i, file_path in enumerate(unprocessed):
        try:
            print(f"📄 Processing {i+1}/{len(unprocessed)}: {file_path.name}")
            
            file_hash = compute_sha256(file_path)
            if not file_hash:
                print(f"   ❌ Failed to compute hash for {file_path.name}")
                continue
            
            # Extract text
            text = extract_text(file_path)
            if not text or len(text.strip()) < 20:
                print(f"   ⚠️  No meaningful text from {file_path.name}")
                continue
            
            # Chunk text
            chunks = chunk_text(text)
            if not chunks:
                print(f"   ⚠️  No chunks generated from {file_path.name}")
                continue
            
            print(f"   ✅ Generated {len(chunks)} chunks from {file_path.name}")
            
            # Generate embeddings
            embeddings = []
            for j in range(0, len(chunks), 32):  # Smaller batch size
                batch = chunks[j:j+32]
                try:
                    embs = model.encode(batch, show_progress_bar=False, convert_to_numpy=True, normalize_embeddings=True)
                    embeddings.append(embs.astype('float32'))
                except Exception as e:
                    print(f"   ❌ Failed to embed batch for {file_path.name}: {e}")
                    continue
            
            if embeddings:
                vectors = np.concatenate(embeddings)
                index.add(vectors)
                
                # Add metadata
                timestamp = datetime.utcnow().isoformat()
                for idx in range(len(chunks)):
                    metadata.append({
                        "file_name": str(file_path.relative_to(KB_DIR)),
                        "file_sha256": file_hash,
                        "chunk_index": idx,
                        "ingest_timestamp": timestamp,
                    })
                
                total_chunks += len(chunks)
                processed += 1
                
                # Update manifest
                with open(MANIFEST_PATH, 'a', encoding='utf-8') as f:
                    entry = {
                        "file_name": str(file_path.relative_to(KB_DIR)),
                        "file_sha256": file_hash,
                        "chunk_count": len(chunks),
                        "ingest_timestamp": timestamp,
                        "model_name": "all-MiniLM-L6-v2",
                        "index_version": version,
                    }
                    f.write(json.dumps(entry) + "\n")
                
                # Save progress every 10 files
                if processed % 10 == 0:
                    faiss_path = IDX_DIR / f"index_v{version}.faiss"
                    pkl_path = IDX_DIR / f"index_v{version}.pkl"
                    
                    faiss.write_index(index, str(faiss_path))
                    with open(pkl_path, 'w') as f:
                        json.dump(metadata, f)
                    
                    elapsed = time.time() - start_time
                    rate = processed / elapsed if elapsed > 0 else 0
                    print(f"   💾 Saved checkpoint: {processed} files, {total_chunks} chunks ({rate:.1f} files/sec)")
            
        except Exception as e:
            print(f"   ❌ Failed to process {file_path.name}: {e}")
            continue
    
    # Final save
    faiss_path = IDX_DIR / f"index_v{version}.faiss"
    pkl_path = IDX_DIR / f"index_v{version}.pkl"
    
    faiss.write_index(index, str(faiss_path))
    with open(pkl_path, 'w') as f:
        json.dump(metadata, f)
    
    print("\n" + "=" * 70)
    print("🎉 ENHANCED INGESTION COMPLETE!")
    print("=" * 70)
    print(f"✅ Processed: {processed} new files")
    print(f"📊 Total chunks: {total_chunks}")
    print(f"📁 Index size: {faiss_path.stat().st_size / 1024 / 1024:.1f}MB")
    print(f"📈 Coverage improvement: +{processed} files")
    
    # Calculate coverage percentage
    total_files_in_kb = len(list(KB_DIR.rglob("*")))
    total_processed = len(existing_hashes) + processed
    coverage = (total_processed / total_files_in_kb) * 100 if total_files_in_kb > 0 else 0
    print(f"🎯 Estimated coverage: {coverage:.1f}%")

if __name__ == "__main__":
    main()
